"""Fill the SDD template from an Extracted object + a rendered diagram PNG.

The template at `templates/Automation_SDD_template.docx` is tokenized with
`{{name}}` placeholders. See `prompts/template_tokens.md` for the full token
list. Three kinds of fill operations happen here:

  1. **Repeating rows.** Tables whose data row contains `{{prefix.field}}`
     tokens (`app.*`, `err.*`, `report.*`) get the row cloned once per item.
  2. **Anchor paragraphs.** A paragraph containing `{{applications_diagram}}`
     is replaced by an inline PNG. A paragraph containing `{{steps}}` is
     replaced by one block of paragraphs per step.
  3. **Scalar tokens.** Everything else is a straight `{{name}}` → string
     replacement, with `[TBD - <reason>]` for missing values.
"""

from __future__ import annotations

import os
import re
from collections.abc import Iterable, Iterator
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from .models import Extracted, Step

TOKEN_RE = re.compile(r"\{\{([a-zA-Z0-9_.]+)\}\}")
_SENTENCE_SPLIT = re.compile(r"(?<=[.!?])\s+")


def _tbd(reason: str = "missing") -> str:
    return f"[TBD - {reason}]"


def fill_template(
    extracted: Extracted,
    diagram_png_path: Path | None,
    out_path: Path,
    template_path: Path | None = None,
) -> Path:
    template_path = template_path or Path(
        os.environ.get("TEMPLATE_PATH", "templates/Automation_SDD_template.docx")
    )
    doc = Document(str(template_path))

    # Repeating rows first — they add new XML and shift indices.
    for table in doc.tables:
        _fill_repeating_rows(table, "app", extracted.applications)
        _fill_repeating_rows(table, "err", extracted.known_errors)
        _fill_repeating_rows(table, "report", extracted.reports)

    # Steps block — design improvements render inline as per-step sub-bullets.
    steps_anchor = _find_paragraph_with_token(doc, "steps")
    if steps_anchor is not None:
        _render_steps_block(steps_anchor, extracted.steps)

    # Diagram embed.
    diagram_anchor = _find_paragraph_with_token(doc, "applications_diagram")
    if diagram_anchor is not None:
        if diagram_png_path is not None and Path(diagram_png_path).is_file():
            _embed_diagram(diagram_anchor, Path(diagram_png_path))
        else:
            _set_paragraph_text(diagram_anchor, _tbd("diagram PNG not provided"))

    # Scalar tokens last.
    scalars = _build_scalar_values(extracted)
    for para in _iter_all_paragraphs(doc):
        _replace_simple_tokens(para, scalars)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _build_scalar_values(ex: Extracted) -> dict[str, str]:
    def s(val, reason: str = "missing") -> str:
        if val is None or val == "":
            return _tbd(reason)
        if isinstance(val, list):
            return "; ".join(str(x) for x in val) if val else _tbd(reason)
        return str(val)

    def s_optional_list(val: list[str]) -> str:
        """For fields where an empty list deliberately means 'none required'
        (document_processing, artificial_intelligence). Both the extractor
        and the narrative prompt default these to ``[]`` when no OCR / AI
        capability is needed — so empty should render as 'None required',
        not [TBD]."""
        return "; ".join(str(x) for x in val) if val else "None required"

    return {
        "project_name": s(ex.project_name, "project name missing"),
        "summary": s(ex.summary, "summary missing"),
        "business_owner": s(ex.business_owner, "owner missing"),
        "jira_project": _tbd("populate manually"),
        "automation_tools": _tbd("populate manually"),
        "btp_services": _tbd("populate manually"),
        "document_processing": s_optional_list(ex.document_processing),
        "new_sdks_objects": _tbd("populate manually"),
        "artificial_intelligence": s_optional_list(ex.artificial_intelligence),
        "credential_management": s(ex.credential_management),
        "tool_selection_rationale": _tbd("populate manually"),
        "business_criticality": s(ex.business_criticality),
        "complexity_score": s(ex.complexity_score),
        "accepted_failure_threshold": s(ex.accepted_failure_threshold),
        "rerun_on_failure": s(ex.rerun_on_failure),
        "schedule_frequency": s(ex.schedule_frequency),
        "bot_utilization_pct": s(ex.bot_utilization_pct),
        "triggers": s(ex.triggers),
    }


def _iter_all_paragraphs(doc) -> Iterator[Paragraph]:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _paragraph_text(para: Paragraph) -> str:
    return "".join(r.text for r in para.runs)


def _set_paragraph_text(para: Paragraph, text: str) -> None:
    for r in para.runs[1:]:
        r._element.getparent().remove(r._element)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _replace_simple_tokens(para: Paragraph, values: dict[str, str]) -> None:
    text = _paragraph_text(para)
    if "{{" not in text:
        return
    new_text = TOKEN_RE.sub(lambda m: values.get(m.group(1), m.group(0)), text)
    if new_text == text:
        return
    _set_paragraph_text(para, new_text)


def _find_paragraph_with_token(doc, token_name: str) -> Paragraph | None:
    target = "{{" + token_name + "}}"
    for para in _iter_all_paragraphs(doc):
        if target in _paragraph_text(para):
            return para
    return None


def _fill_repeating_rows(table, prefix: str, items: Iterable) -> None:
    pattern = re.compile(r"\{\{" + re.escape(prefix) + r"\.[a-zA-Z0-9_]+\}\}")
    template_rows = [
        row for row in table.rows if pattern.search(" ".join(c.text for c in row.cells))
    ]
    if not template_rows:
        return

    proto_xml = template_rows[0]._element
    parent = proto_xml.getparent()
    proto_index = list(parent).index(proto_xml)

    items_list = list(items)
    new_row_xmls = []
    for item in items_list:
        cloned = deepcopy(proto_xml)
        for t in cloned.iter(qn("w:t")):
            if t.text and "{{" in t.text:
                t.text = TOKEN_RE.sub(
                    lambda m, _item=item: _resolve_item_token(m.group(1), prefix, _item),
                    t.text,
                )
        new_row_xmls.append(cloned)

    for offset, nr in enumerate(new_row_xmls):
        parent.insert(proto_index + 1 + offset, nr)

    for row in template_rows:
        parent.remove(row._element)


def _resolve_item_token(token: str, prefix: str, item) -> str:
    expected = prefix + "."
    if not token.startswith(expected):
        return "{{" + token + "}}"
    field = token[len(expected) :]
    val = getattr(item, field, None)
    if val is None or val == "":
        return _tbd(f"{prefix}.{field}")
    if isinstance(val, list):
        return "; ".join(str(x) for x in val) if val else _tbd(f"{prefix}.{field}")
    return str(val)


def _embed_diagram(anchor: Paragraph, png_path: Path) -> None:
    for r in anchor.runs:
        r.text = ""
    if anchor.runs:
        for r in anchor.runs[1:]:
            r._element.getparent().remove(r._element)
        run = anchor.runs[0]
    else:
        run = anchor.add_run()
    run.add_picture(str(png_path), width=Inches(6.5))


def _render_steps_block(anchor: Paragraph, steps: list[Step]) -> None:
    """Render the SBS flow as grouped step blocks, followed by a separate
    "Developer notes to consider" section if any step has a design note.

    Layout per step: group header (bold) when the group changes → step header
    ``N) summary`` with N) bold, no bullet → bulleted sub-items in order:
    action_detail lines (the click-by-click / API-call prose), then decision
    rules split sentence-by-sentence, then exception paths.
    """
    if not steps:
        _set_paragraph_text(anchor, _tbd("no steps extracted"))
        return

    paragraphs: list[list[tuple[bool, str]] | None] = []
    previous_group: str | None = None

    for step in steps:
        if step.group and step.group != previous_group:
            if previous_group is not None:
                paragraphs.append(None)
            paragraphs.append([(True, step.group)])
            previous_group = step.group

        paragraphs.append(
            [
                (True, f"{step.number})"),
                (False, f" {step.summary}"),
            ]
        )

        for line in _iter_action_lines(step.action_detail):
            paragraphs.append([(False, f"    •  {line}")])
        for sentence in _split_sentences(step.decision_logic):
            paragraphs.append([(False, f"    •  {sentence}")])
        for exc in step.exception_paths:
            paragraphs.append([(False, f"    •  {exc}")])

    notes = [(s.number, s.design_note) for s in steps if s.design_note]
    if notes:
        paragraphs.append(None)
        paragraphs.append([(True, "Developer notes to consider")])
        for number, note in notes:
            paragraphs.append([(False, f"    •  Step {number}: {note}")])

    _render_run_paragraphs(anchor, paragraphs)


def _iter_action_lines(action_detail: str) -> list[str]:
    """Action detail is freeform prose. Split first on hard newlines (the
    LLM sometimes returns a list); for any line that still reads as multiple
    sentences, split sentence-by-sentence so each click / call gets its own
    bullet."""
    if not action_detail or not action_detail.strip():
        return []
    out: list[str] = []
    for hard_line in action_detail.splitlines():
        hard_line = hard_line.strip().lstrip("-•* ").strip()
        if not hard_line:
            continue
        for sentence in _split_sentences(hard_line) or [hard_line]:
            out.append(sentence)
    return out


def _split_sentences(text: str) -> list[str]:
    """Split a string into sentences on `.`/`!`/`?` boundaries, preserving
    the trailing punctuation. Empty or whitespace-only input → empty list."""
    if not text or not text.strip():
        return []
    parts = _SENTENCE_SPLIT.split(text.strip())
    return [p.strip() for p in parts if p.strip()]


def _render_run_paragraphs(
    anchor: Paragraph, paragraphs: list[list[tuple[bool, str]] | None]
) -> None:
    """Replace the anchor paragraph + insert following paragraphs to match
    the given run-fragment list. Each item is either ``None`` (blank line)
    or a list of ``(bold, text)`` tuples that become runs in order."""
    if not paragraphs:
        _set_paragraph_text(anchor, "")
        return

    _set_paragraph_runs(anchor, paragraphs[0])
    cursor_xml = anchor._element
    for runs in paragraphs[1:]:
        new_p = deepcopy(cursor_xml)
        for child in list(new_p):
            if child.tag == qn("w:r"):
                new_p.remove(child)
        if runs:
            for bold, text in runs:
                new_p.append(_make_run(text, bold=bold))
        cursor_xml.addnext(new_p)
        cursor_xml = new_p


def _set_paragraph_runs(para: Paragraph, runs: list[tuple[bool, str]] | None) -> None:
    """Clear all runs in ``para`` and replace them with the given fragments."""
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    if not runs:
        return
    p_element = para._element
    for bold, text in runs:
        p_element.append(_make_run(text, bold=bold))


def _make_run(text: str, *, bold: bool = False) -> OxmlElement:
    r = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rpr.append(b)
        r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r
